import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋_GB2312"
SUBTITLE_FONT = "楷体_GB2312"
LEVEL1_FONT = "黑体"

TITLE_SIZE = Pt(22)  # 2号
BODY_SIZE = Pt(16)   # 3号
PAGE_NUM_SIZE = Pt(14)  # 4号

LINE_SPACING = Pt(29)
TWO_CHAR_INDENT = Pt(32)

DEFAULT_DATA = {
    "title": "关于开展年度工作总结的通知",
    "recipients": ["各相关单位"],
    "body": [
        "为全面总结年度工作成果，梳理经验做法，现就有关事项通知如下。",
        "一、总体要求",
        "（一）突出重点。各单位要围绕中心任务，突出亮点工作。",
        "1. 做到数据准确、材料完整。",
        "（1）按时报送，总结材料不超过三页。",
        "请于2月15日前报送电子版材料。",
    ],
    "attachments": ["年度工作总结模板"],
}

MARKDOWN_EXTENSIONS = {".md", ".markdown"}

HEADING_STYLE_MAP = {
    1: {"font": TITLE_FONT, "size": TITLE_SIZE, "align": WD_ALIGN_PARAGRAPH.CENTER, "indent": None},
    2: {"font": LEVEL1_FONT, "size": BODY_SIZE, "align": WD_ALIGN_PARAGRAPH.LEFT, "indent": TWO_CHAR_INDENT},
    3: {"font": SUBTITLE_FONT, "size": BODY_SIZE, "align": WD_ALIGN_PARAGRAPH.LEFT, "indent": TWO_CHAR_INDENT},
    4: {"font": BODY_FONT, "size": BODY_SIZE, "align": WD_ALIGN_PARAGRAPH.LEFT, "indent": TWO_CHAR_INDENT},
    5: {"font": BODY_FONT, "size": BODY_SIZE, "align": WD_ALIGN_PARAGRAPH.LEFT, "indent": TWO_CHAR_INDENT},
}


def _set_run_font(run, font_name, size, *, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    run.bold = bold


def _set_paragraph_format(paragraph, *, first_line_indent=None, left_indent=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    if first_line_indent is not None:
        fmt.first_line_indent = first_line_indent
    if left_indent is not None:
        fmt.left_indent = left_indent


def _normalize_quotes(text: str) -> str:
    if not text:
        return text
    return re.sub(r'"([^"\n]+)"', r"“\1”", text)


def _add_text_paragraph(document, text, font_name, size, *, align=None,
                        first_line_indent=None, left_indent=None):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    _set_paragraph_format(
        paragraph,
        first_line_indent=first_line_indent,
        left_indent=left_indent,
    )
    run = paragraph.add_run(_normalize_quotes(text))
    _set_run_font(run, font_name, size, bold=(font_name == SUBTITLE_FONT))
    return paragraph


def _add_blank_paragraph(document):
    paragraph = document.add_paragraph()
    _set_paragraph_format(paragraph)
    return paragraph


def _add_footer_page_number(document):
    section = document.sections[0]
    footer = section.footer
    if not footer.paragraphs:
        paragraph = footer.add_paragraph()
    else:
        paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        paragraph._p.remove(run._r)
    _set_paragraph_format(paragraph)

    run_left = paragraph.add_run("— ")
    _set_run_font(run_left, "宋体", PAGE_NUM_SIZE)
    run_field = paragraph.add_run()
    _set_run_font(run_field, "宋体", PAGE_NUM_SIZE)
    r = run_field._r
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r.append(fld_char_begin)
    r.append(instr_text)
    r.append(fld_char_sep)
    r.append(fld_char_end)

    run_right = paragraph.add_run(" —")
    _set_run_font(run_right, "宋体", PAGE_NUM_SIZE)


def _choose_font_for_paragraph(text):
    stripped = text.strip()
    if stripped.startswith(tuple("一二三四五六七八九十")) and "、" in stripped[:3]:
        return LEVEL1_FONT
    if stripped.startswith("（") and "）" in stripped[:4]:
        return SUBTITLE_FONT
    return BODY_FONT


def _normalize_recipients(value):
    if value is None:
        return ""
    if isinstance(value, list):
        return "、".join([str(item).strip() for item in value if str(item).strip()])
    return str(value).strip()


def _normalize_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    return [text] if text else []


def _parse_body_text(text):
    if not text:
        return []
    lines = []
    for line in text.splitlines():
        cleaned = line.strip()
        if cleaned:
            lines.append(cleaned)
    return lines


def _load_json_input(path):
    if path == "-":
        content = sys.stdin.read()
    else:
        content = Path(path).read_text(encoding="utf-8")
    return json.loads(content)


def _read_text_input(path: str) -> str:
    if path == "-":
        return sys.stdin.read()
    return Path(path).read_text(encoding="utf-8")


def _parse_simple_front_matter(text: str):
    """解析受控 YAML Front Matter（不依赖第三方库）。

    仅支持：
    - key: value（字符串）
    - key: （空值后跟列表）
      - item
      - item
    - key: [a, b]（简易行内列表）

    返回：(meta: dict, rest_text: str)
    """
    stripped = text.lstrip("\ufeff")
    if not stripped.startswith("---\n") and stripped != "---":
        return {}, text

    lines = stripped.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, text

    meta = {}
    i = 1
    current_list_key = None
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            i += 1
            break
        if not line.strip() or line.lstrip().startswith("#"):
            i += 1
            continue

        if line.startswith("  - ") or line.startswith("- "):
            if current_list_key:
                item = line.split("-", 1)[1].strip()
                if item:
                    meta.setdefault(current_list_key, []).append(item)
            i += 1
            continue

        current_list_key = None
        if ":" in line:
            key, value = line.split(":", 1)
            key = key.strip()
            value = value.strip()
            if not key:
                i += 1
                continue
            if value == "":
                meta[key] = []
                current_list_key = key
                i += 1
                continue
            if value.startswith("[") and value.endswith("]"):
                inner = value[1:-1].strip()
                if inner:
                    meta[key] = [part.strip().strip("\"'") for part in inner.split(",") if part.strip()]
                else:
                    meta[key] = []
                i += 1
                continue
            meta[key] = value.strip("\"'")
        i += 1

    rest_text = "\n".join(lines[i:])
    return meta, rest_text


def _parse_controlled_markdown(text: str):
    """将受控 Markdown 解析为 meta + blocks。

    规则：
    - 每一行非空文本 = 一个段落（自然换行即自然段）
    - 仅识别行首连续 # 的数量（1~5）作为标题层级
    - 行中出现的 # 一律不处理
    """
    meta, body_text = _parse_simple_front_matter(text)

    blocks = []
    for raw_line in body_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("#"):
            level = 0
            for ch in line:
                if ch == "#":
                    level += 1
                else:
                    break
            if 1 <= level <= 5:
                content = line[level:].lstrip()
                if content:
                    blocks.append({"type": f"h{level}", "text": content})
                    continue
        blocks.append({"type": "p", "text": line.strip()})

    # 从 blocks 中提取第一条 h1 作为 title（同时保留 blocks 供渲染使用）
    title = ""
    for b in blocks:
        if b["type"] == "h1":
            title = b["text"].strip()
            break

    data = {
        "title": title,
        "recipients": meta.get("recipients"),
        "attachments": meta.get("attachments"),
        "signer": meta.get("signer"),
        "date": meta.get("date"),
        "blocks": blocks,
    }
    return data


def _load_markdown_input(path: str):
    content = _read_text_input(path)
    return _parse_controlled_markdown(content)


def build_document(data, output_path: Path):
    title = str(data.get("title", "")).strip()
    recipients = _normalize_recipients(data.get("recipients"))
    body = data.get("body", [])
    attachments = _normalize_list(data.get("attachments"))
    signer = str(data.get("signer", "")).strip()
    date = str(data.get("date", "")).strip()
    blocks = data.get("blocks")

    if isinstance(body, str):
        body = _parse_body_text(body)
    elif isinstance(body, list):
        body = [str(item).strip() for item in body if str(item).strip()]
    else:
        body = []

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    if isinstance(blocks, list) and blocks:
        title_rendered = False
        first_h1_consumed = False

        # 1) 标题：只渲染第一条 h1
        for block in blocks:
            if block.get("type") == "h1":
                text = str(block.get("text", "")).strip()
                if text:
                    _add_text_paragraph(
                        document,
                        text,
                        TITLE_FONT,
                        TITLE_SIZE,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                    )
                    _add_blank_paragraph(document)
                    title_rendered = True
                first_h1_consumed = True
                break

        # 若未提供 h1，则回退使用 data["title"]（兼容）
        if not title_rendered and title:
            _add_text_paragraph(
                document,
                title,
                TITLE_FONT,
                TITLE_SIZE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _add_blank_paragraph(document)
            title_rendered = True

        # 2) 主送机关（Front Matter 提供）
        if recipients:
            _add_text_paragraph(
                document,
                f"{recipients}：",
                BODY_FONT,
                BODY_SIZE,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )

        # 3) 正文 blocks：跳过第一条 h1
        skipped_first_h1 = False
        for block in blocks:
            btype = block.get("type")
            text = str(block.get("text", "")).strip()
            if not text:
                continue

            if btype == "h1" and not skipped_first_h1:
                skipped_first_h1 = True
                continue

            if btype and btype.startswith("h"):
                try:
                    level = int(btype[1:])
                except ValueError:
                    level = 0
                style = HEADING_STYLE_MAP.get(level)
                if style and level != 1:
                    _add_text_paragraph(
                        document,
                        text,
                        style["font"],
                        style["size"],
                        align=style["align"],
                        first_line_indent=style["indent"],
                    )
                else:
                    _add_text_paragraph(
                        document,
                        text,
                        BODY_FONT,
                        BODY_SIZE,
                        first_line_indent=TWO_CHAR_INDENT,
                    )
            else:
                _add_text_paragraph(
                    document,
                    text,
                    BODY_FONT,
                    BODY_SIZE,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_line_indent=TWO_CHAR_INDENT,
                )

    else:
        if title:
            _add_text_paragraph(
                document,
                title,
                TITLE_FONT,
                TITLE_SIZE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _add_blank_paragraph(document)

        if recipients:
            _add_text_paragraph(
                document,
                f"{recipients}：",
                BODY_FONT,
                BODY_SIZE,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )

        for paragraph_text in body:
            font_name = _choose_font_for_paragraph(paragraph_text)
            _add_text_paragraph(
                document,
                paragraph_text,
                font_name,
                BODY_SIZE,
                align=(
                    WD_ALIGN_PARAGRAPH.JUSTIFY
                    if font_name == BODY_FONT
                    else WD_ALIGN_PARAGRAPH.LEFT
                ),
                first_line_indent=TWO_CHAR_INDENT,
            )

    if attachments:
        _add_blank_paragraph(document)
        if len(attachments) == 1:
            _add_text_paragraph(
                document,
                f"附件：{attachments[0]}",
                BODY_FONT,
                BODY_SIZE,
                left_indent=TWO_CHAR_INDENT,
            )
        else:
            for idx, name in enumerate(attachments, 1):
                prefix = "附件：" if idx == 1 else " " * 4
                _add_text_paragraph(
                    document,
                    f"{prefix}{idx}. {name}",
                    BODY_FONT,
                    BODY_SIZE,
                    left_indent=TWO_CHAR_INDENT,
                )

    if signer or date:
        _add_blank_paragraph(document)
        if signer:
            _add_text_paragraph(
                document,
                signer,
                BODY_FONT,
                BODY_SIZE,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )
        if date:
            _add_text_paragraph(
                document,
                date,
                BODY_FONT,
                BODY_SIZE,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )

    _add_footer_page_number(document)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="生成符合公文格式的Word文档")
    parser.add_argument("--input", help="输入JSON文件路径，或 '-' 从STDIN读取")
    parser.add_argument("--md", help="输入Markdown文件路径（受控格式），或 '-' 从STDIN读取")
    parser.add_argument("--title", help="标题")
    parser.add_argument("--recipients", help="主送机关，多个用顿号分隔")
    parser.add_argument("--body", help="正文，多段用换行分隔")
    parser.add_argument("--body-file", help="正文文本文件，每行一段")
    parser.add_argument("--attachment", action="append", help="附件名称，可多次传入")
    parser.add_argument(
        "-o",
        "--output",
        default="示例公文.docx",
        help="输出文件路径（默认：示例公文.docx）",
    )
    args = parser.parse_args()

    data = {}
    has_custom = any([
        args.title,
        args.recipients,
        args.body,
        args.body_file,
        args.attachment,
    ])

    if args.md:
        data = _load_markdown_input(args.md)
    elif args.input:
        suffix = Path(args.input).suffix.lower() if args.input != "-" else ""
        if suffix in MARKDOWN_EXTENSIONS:
            data = _load_markdown_input(args.input)
        else:
            data = _load_json_input(args.input)
    elif has_custom:
        data = {}
    else:
        data = DEFAULT_DATA.copy()

    if args.title:
        data["title"] = args.title
    if args.recipients:
        data["recipients"] = args.recipients

    if args.body_file:
        body_text = Path(args.body_file).read_text(encoding="utf-8")
        data["body"] = _parse_body_text(body_text)
    elif args.body:
        data["body"] = _parse_body_text(args.body)

    if args.attachment:
        data["attachments"] = args.attachment

    output_path = Path(args.output).resolve()
    build_document(data, output_path)
    print(f"已生成：{output_path}")


if __name__ == "__main__":
    main()
